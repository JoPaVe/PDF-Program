#from PyPDF2 import PdfFileReader
import fitz
import os
from docx import Document
from typing import List, Tuple  
from tkinter import Tk, Label
from tkinter.filedialog import asksaveasfilename,askdirectory
from PIL import ImageTk, Image  

# Extrahiert aus Highlight-Annotation den Text -> Gibt String des Highlights zurück
def _parse_highlight(annot: fitz.Annot, wordlist: List[Tuple[float, float, float, float, str, int, int, int]]) -> str:
    points = annot.vertices
    quad_count = int(len(points) / 4)
    sentences = []
    for i in range(quad_count):
        # where the highlighted part is
        r = fitz.Quad(points[i * 4 : i * 4 + 4]).rect

        words = [w for w in wordlist if fitz.Rect(w[:4]).intersects(r)] #Überschneidung von Rectangle von Highlight und Gesamttext
        sentences.append(" ".join(w[4] for w in words))
    sentence = " ".join(sentences)
    return sentence


def handle_page(page) -> list:
    wordlist = page.get_text("words")  # list of words on page
    wordlist.sort(key=lambda w: (w[3], w[0]))  # ascending y, then x

    highlights = []
    annot = page.firstAnnot
    while annot:
        if annot.type[0] == 8: #Extrahier Highlights
            highlights.append((page.number,"Highlight",_parse_highlight(annot, wordlist), annot.info["title"]))
        elif annot.type[0] == 0: #Extrahiere Texts
            highlights.append((page.number,"Text",annot.info["content"], annot.info["title"]))
        annot = annot.next
    return highlights

def export_extract(text: list, currentdoc: str, outdoc):

    docname = currentdoc.split("/")[-1].split(".")[0] #Nehme Name aus Path
    outdoc.add_heading(docname)
    
    for paras in text: #Erstelle Output im Doc
        pagenum = paras[0]
        pagenum += 1
        paratype = paras[1]
        paratext = paras[2]
        paratitle = paras[3]

        if paratype == "Highlight":
            outdoc.add_paragraph(f'{paratype} von {paratitle}: S. {str(pagenum)} "{paratext}"')
        elif paratype == "Text":
            outdoc.add_paragraph(f'{paratype} von {paratitle}: S. {str(pagenum)} {paratext}')
    return outdoc

def extract_files(filepaths: list, preview_box = None, open_doc_directly = False):
    
    document = Document()
    for docs in filepaths:
        doc = fitz.open(docs)

        highlights = []
        for page in doc:
            highlights += handle_page(page)
        if preview_box == None: #Wenn nicht für Preview_box extrahiert wird, soll Dokument ausgegeben werden
            document = export_extract(highlights, docs, document)
    
    if preview_box != None: # Wenn für Preview_box extrahiert wird
        if len(highlights) == 0:
            preview_box.insert(1.0,"Keine Kommentare gefunden! ")
        for text in highlights:
            preview_box.insert(1.0,text)
            preview_box.insert(1.0,"\n\n")
    else:
        Tk().withdraw()
        path = asksaveasfilename(filetypes = [('Word','.docx')],defaultextension = [('Word','.docx')],title="Ausgabepfad des Markierungsdokuments")
        #outdocpath = os.path.join(path,"MarkierungenderDokumente.docx")
        document.save(path) 
        if open_doc_directly == True:
            os.system('start '+str(path))

